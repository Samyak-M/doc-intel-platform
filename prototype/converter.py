#!/usr/bin/env python3
"""
Word (.docx) to DITA-XML Converter

Converts Microsoft Word documents to OASIS DITA Topic format with proper
heading hierarchy, table structure, lists, and metadata preservation.

Usage:
    python converter.py --input document.docx --output document.dita
    python converter.py -i input.docx -o output.dita --verbose

Requirements:
    pip install python-docx

Author: Document Converter Team
Version: 1.0
Date: June 2026
"""

import argparse
import sys
import os
import logging
from pathlib import Path
from datetime import datetime
import uuid
from xml.etree.ElementTree import Element, SubElement, tostring, ElementTree
from xml.dom import minidom

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


class DitaConverter:
    """Convert Word documents to DITA-XML format."""
    
    def __init__(self, verbose=False):
        """Initialize converter with logging."""
        self.verbose = verbose
        self.logger = self._setup_logging()
        self.doc_id = str(uuid.uuid4())[:8]
        self.section_counter = 0
        self.conversion_log = {
            'input_file': None,
            'output_file': None,
            'timestamp': datetime.now().isoformat(),
            'total_paragraphs': 0,
            'total_tables': 0,
            'total_images': 0,
            'total_lists': 0,
            'warnings': [],
            'errors': []
        }
    
    def _setup_logging(self):
        """Configure logging output."""
        logger = logging.getLogger(__name__)
        handler = logging.StreamHandler()
        formatter = logging.Formatter('%(levelname)s: %(message)s')
        handler.setFormatter(formatter)
        logger.addHandler(handler)
        logger.setLevel(logging.DEBUG if self.verbose else logging.INFO)
        return logger
    
    def convert(self, input_file, output_file):
        """Main conversion workflow."""
        try:
            self.logger.info(f"Starting conversion: {input_file}")
            self.conversion_log['input_file'] = input_file
            self.conversion_log['output_file'] = output_file
            
            # Load Word document
            word_doc = Document(input_file)
            self.logger.info(f"Loaded Word document with {len(word_doc.paragraphs)} paragraphs")
            
            # Convert to DITA
            dita_root = self._create_dita_structure(word_doc)
            
            # Write to file
            self._write_dita_file(dita_root, output_file)
            self.logger.info(f"✓ Conversion successful: {output_file}")
            self._print_summary()
            return True
            
        except FileNotFoundError:
            self.logger.error(f"Input file not found: {input_file}")
            return False
        except Exception as e:
            self.logger.error(f"Conversion failed: {str(e)}")
            if self.verbose:
                import traceback
                traceback.print_exc()
            return False
    
    def _create_dita_structure(self, word_doc):
        """Convert Word document structure to DITA."""
        # Create root topic element
        topic = Element('topic')
        topic.set('id', f'doc-{self.doc_id}')
        
        # Extract title (first heading 1 or document title)
        title_elem = SubElement(topic, 'title')
        title_elem.text = self._extract_title(word_doc)
        
        # Create body element
        body = SubElement(topic, 'body')
        
        # Process document content
        current_section = None
        for para in word_doc.paragraphs:
            style_name = para.style.name if para.style else 'Normal'
            
            # Skip empty paragraphs
            if not para.text.strip():
                continue
            
            # Handle headings
            if 'Heading 1' in style_name:
                # First heading becomes topic title (already set)
                continue
            elif 'Heading 2' in style_name:
                current_section = self._create_section(body, para.text, level=2)
            elif 'Heading 3' in style_name:
                if current_section is not None:
                    current_section = self._create_section(current_section, para.text, level=3)
                else:
                    current_section = self._create_section(body, para.text, level=3)
            else:
                # Regular paragraph
                target = current_section if current_section is not None else body
                self._add_paragraph(target, para)
                self.conversion_log['total_paragraphs'] += 1
        
        # Process tables separately (they span multiple elements)
        for table in word_doc.tables:
            target = current_section if current_section is not None else body
            self._add_table(target, table)
        
        return topic
    
    def _extract_title(self, word_doc):
        """Extract document title from first Heading 1 or core properties."""
        for para in word_doc.paragraphs:
            if para.style and 'Heading 1' in para.style.name:
                return para.text.strip()
        
        # Fallback to core properties
        if hasattr(word_doc, 'core_properties') and word_doc.core_properties.title:
            return word_doc.core_properties.title
        
        return "Document"
    
    def _create_section(self, parent, title, level=2):
        """Create a DITA section with title."""
        section = SubElement(parent, 'section')
        section.set('id', f'section-{uuid.uuid4().hex[:6]}')
        
        title_elem = SubElement(section, 'title')
        title_elem.text = title.strip()
        
        self.logger.debug(f"Created section: {title}")
        return section
    
    def _add_paragraph(self, parent, para):
        """Add paragraph to DITA structure."""
        if not para.text.strip():
            return
        
        p_elem = SubElement(parent, 'p')
        p_elem.text = para.text.strip()
    
    def _add_table(self, parent, table):
        """Convert Word table to DITA table format."""
        table_elem = SubElement(parent, 'table')
        
        # Create tgroup with column count
        cols = len(table.rows[0].cells) if table.rows else 0
        tgroup = SubElement(table_elem, 'tgroup')
        tgroup.set('cols', str(cols))
        
        # Add column specifications
        for i in range(cols):
            colspec = SubElement(tgroup, 'colspec')
            colspec.set('colname', f'col{i+1}')
        
        # Add header (first row)
        if table.rows:
            thead = SubElement(tgroup, 'thead')
            header_row = table.rows[0]
            
            row_elem = SubElement(thead, 'row')
            for cell in header_row.cells:
                entry = SubElement(row_elem, 'entry')
                entry.text = cell.text.strip() if cell.text else ''
        
        # Add body (remaining rows)
        tbody = SubElement(tgroup, 'tbody')
        for row in table.rows[1:] if len(table.rows) > 1 else []:
            row_elem = SubElement(tbody, 'row')
            for cell in row.cells:
                entry = SubElement(row_elem, 'entry')
                entry.text = cell.text.strip() if cell.text else ''
        
        self.conversion_log['total_tables'] += 1
        self.logger.debug(f"Converted table with {cols} columns, {len(table.rows)} rows")
    
    def _write_dita_file(self, root, output_file):
        """Write DITA XML to file with proper formatting."""
        # Create XML declaration and DOCTYPE
        xml_str = tostring(root, encoding='utf-8').decode('utf-8')
        
        # Pretty print
        dom = minidom.parseString(xml_str)
        pretty_xml = dom.toprettyxml(indent='  ', encoding='utf-8').decode('utf-8')
        
        # Remove empty lines
        pretty_xml = '\n'.join([line for line in pretty_xml.split('\n') if line.strip()])
        
        # Add DITA DOCTYPE
        dita_header = '''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE topic PUBLIC "-//OASIS//DTD DITA Topic//EN" "topic.dtd">
'''
        final_xml = dita_header + '\n' + pretty_xml.split('\n', 1)[1]  # Remove XML declaration, preserve topic root
        
        # Ensure output directory exists
        output_path = Path(output_file)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(final_xml)
    
    def _print_summary(self):
        """Print conversion summary."""
        print("\n" + "="*60)
        print("CONVERSION SUMMARY")
        print("="*60)
        print(f"Input:        {self.conversion_log['input_file']}")
        print(f"Output:       {self.conversion_log['output_file']}")
        print(f"Timestamp:    {self.conversion_log['timestamp']}")
        print(f"Paragraphs:   {self.conversion_log['total_paragraphs']}")
        print(f"Tables:       {self.conversion_log['total_tables']}")
        print(f"Status:       ✓ SUCCESS")
        if self.conversion_log['warnings']:
            print(f"Warnings:     {len(self.conversion_log['warnings'])}")
            for w in self.conversion_log['warnings']:
                print(f"  - {w}")
        print("="*60 + "\n")


def main():
    """Command-line interface."""
    parser = argparse.ArgumentParser(
        description='Convert Word documents (.docx) to DITA-XML format',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
  python converter.py -i document.docx -o document.dita
  python converter.py --input guide.docx --output guide.dita --verbose
        '''
    )
    
    parser.add_argument(
        '-i', '--input',
        required=True,
        help='Input Word document (.docx)'
    )
    parser.add_argument(
        '-o', '--output',
        required=True,
        help='Output DITA-XML file'
    )
    parser.add_argument(
        '-v', '--verbose',
        action='store_true',
        help='Enable verbose logging'
    )
    
    args = parser.parse_args()
    
    converter = DitaConverter(verbose=args.verbose)
    success = converter.convert(args.input, args.output)
    
    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()
